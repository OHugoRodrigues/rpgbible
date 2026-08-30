from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


SOURCE = Path(r"C:\Users\gugar\Downloads\Peregrino_Hackathon_Produto_e_User_Stories.docx")
OUTPUT = Path(r"C:\Users\gugar\OneDrive\Documentos\ChatGPT\RPJesus\deliverables\Peregrino_Hackathon_Produto_e_User_Stories_Atualizado.docx")


def replace_exact(document: Document, old: str, new: str) -> None:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == old.strip():
            paragraph.text = new
            return
    raise ValueError(f"Paragraph not found: {old[:70]}")


def replace_cell_exact(document: Document, old: str, new: str) -> None:
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() == old.strip():
                    cell.text = new
                    return
    raise ValueError(f"Cell not found: {old[:70]}")


def insert_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    created = Paragraph(new_p, paragraph._parent)
    created.add_run(text)
    return created


def shade_paragraph(paragraph, fill: str, border: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:color"), border)
    borders.append(bottom)
    p_pr.append(borders)


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement("w:keepNext"))


def add_bullet(document: Document, text: str):
    paragraph = document.add_paragraph(text, style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    return paragraph


def add_story(document: Document, story_id: str, title: str, statement: str, criteria: list[str], owner: str) -> None:
    heading = document.add_heading(f"{story_id} - {title}", level=2)
    keep_with_next(heading)
    metadata = document.add_paragraph()
    run = metadata.add_run(f"Prioridade: P0  |  Dono: {owner}")
    run.bold = True
    run.font.color.rgb = RGBColor(33, 86, 115)
    metadata.paragraph_format.space_after = Pt(4)
    paragraph = document.add_paragraph(statement)
    paragraph.paragraph_format.space_after = Pt(5)
    label = document.add_paragraph()
    label.add_run("Critérios de aceite").bold = True
    label.paragraph_format.space_after = Pt(2)
    for criterion in criteria:
        add_bullet(document, criterion)


document = Document(SOURCE)

# Capa e contexto do documento.
document.paragraphs[2].text = "Definição de Produto, Experiência Adaptativa e User Stories"
document.paragraphs[4].text = "Referência consolidada do universo de jornadas, personalização, minigames, progressão, IA segura e colecionáveis."
update_note = insert_after(
    document.paragraphs[4],
    "VERSÃO CONSOLIDADA DO PROTÓTIPO - Atualizada com as decisões de produto, experiência e arquitetura definidas durante o hackathon.",
)
update_note.paragraph_format.space_before = Pt(8)
update_note.paragraph_format.space_after = Pt(12)
update_note.paragraph_format.left_indent = Pt(10)
update_note.paragraph_format.right_indent = Pt(10)
update_note.runs[0].bold = True
update_note.runs[0].font.size = Pt(9)
update_note.runs[0].font.color.rgb = RGBColor(255, 255, 255)
shade_paragraph(update_note, "17324D", "E6B95C")

# Atualizações locais nos requisitos existentes que ficaram desatualizados.
replace_exact(
    document,
    "Peregrino é uma plataforma de engajamento bíblico que transforma as Escrituras em jornadas interativas, ajudando jovens e adolescentes a criar constância, aprender e avançar em sua caminhada. O RPG é a mecânica; a formação bíblica é o produto.",
    "Peregrino é uma plataforma de engajamento bíblico que transforma as Escrituras em jornadas interativas, ajudando jovens e adolescentes a criar constância, aprender e avançar em sua caminhada. O RPG é a mecânica; a formação bíblica é o produto. A experiência é mobile-first, imersiva, visual e adaptativa: menos menus e mais jornada acontecendo em tela cheia.",
)
replace_exact(
    document,
    "3. A Provação — Davi e Golias. Capítulo jogável do hackathon.",
    "3. A Provação — Davi e Golias. Momento central e confronto em AR do piloto.",
)
replace_exact(
    document,
    "Descoberta  •  Minigame  •  Decisão  •  Quiz  •  Interação AR  •  Exploração  •  Reflexão",
    "Diálogo de calibragem → Trilha adaptativa → Contexto interativo → Minigames com tempo → Decisão → AR ou simulação → Consolidação → Recompensa → Constância",
)
replace_exact(document, "8. Backlog do MVP — 36 horas", "8. Backlog do MVP — menos de 24 horas")
replace_exact(
    document,
    "O backlog abaixo prioriza um vertical slice completo. Qualquer item fora deste fluxo deve ser tratado como P1/P2 até o caminho principal funcionar ponta a ponta.",
    "O backlog abaixo prioriza uma demonstração ponta a ponta em menos de 24 horas. Vale mais ver a jornada funcionando com qualidade do que introduzir complexidade técnica sem impacto no pitch. Qualquer item fora do fluxo demonstrável permanece P1/P2.",
)
replace_exact(
    document,
    "Destacar “A Provação” como capítulo jogável do hackathon.",
    "Exibir cinco etapas jogáveis no protótipo, preservando a progressão narrativa até A Provação.",
)
replace_exact(
    document,
    "Demais capítulos podem ficar como preview/bloqueados.",
    "Jornadas futuras permanecem como preview/bloqueadas; a jornada piloto de Davi deve parecer completa.",
)
replace_exact(
    document,
    "Nota de produto: Essa é a principal experiência WOW. O especialista de AR não deve receber outra frente crítica até esse fluxo funcionar.",
    "Nota de produto: Essa é a principal experiência WOW. O fluxo WebXR real deve coexistir com uma simulação de precisão para que a demonstração continue jogável em desktop ou em aparelho incompatível.",
)
replace_exact(
    document,
    "Dev 1 — Front/Experiência\nOnboarding, Peregrino, Jornadas, capítulos, minigames, quiz, recompensas, estante, inventário e constância. Deve trabalhar com mocks enquanto o AR não estiver integrado.",
    "Dev 1 — Front/Experiência\nJornada em tela cheia, criação do Peregrino, calibragem imersiva, mapa, cinco minigames, HUD, Parakletos, recompensas, estante, inventário e constância. Deve manter o fallback jogável enquanto o AR real não estiver disponível.",
)
replace_exact(
    document,
    "Dev 2 — Arquitetura/Back\nModelo mínimo de domínio, persistência, score, progressão e integração. Em caso de risco de prazo, priorizar persistência local/leve no piloto.",
    "Dev 2 — Arquitetura/IA\nModelo de domínio, persistência, score, progressão, perfis adaptativos, endpoint seguro de LLM e fallback determinístico. Chaves e prompts de sistema nunca ficam no cliente; no piloto, o estado persiste localmente.",
)
replace_exact(
    document,
    "Dev 3 — VC/AR\nUma missão principal: fazer Davi x Golias aparecer em superfície e permitir o lançamento da pedra com feedback convincente. É o dono do momento WOW.",
    "Dev 3 — VC/AR\nFazer Davi x Golias aparecer em superfície via WebXR hit-test, permitir posicionamento e lançamento com feedback convincente e garantir uma saída segura da sessão. É o dono do momento WOW.",
)
replace_exact(
    document,
    "Designer\nDesign system mínimo, Peregrino, Home, Jornadas, narrativa, HUD AR, recompensa, estante/inventário e assets estritamente necessários.",
    "Designer\nDireção visual RPG em pixel art, tela inicial cinematográfica, Peregrinos em diferentes estágios, Parakletos em múltiplas poses, mapas, HUD, itens e recompensas. A interface deve competir por atenção com jogos juvenis sem copiar marcas ou padrões proprietários.",
)
replace_exact(
    document,
    "Produto/PO\nFechar conteúdo, referências, regras, critérios de aceite, priorização e decisões. Remover ambiguidade do caminho do time e impedir expansão de escopo.",
    "Produto/PO\nFechar conteúdo, referências, regras, critérios de aceite e guardrails teológicos. Revisar paráfrases e impedir que IA, Parakletos ou mecânicas de recompensa assumam autoridade espiritual ou substituam a Escritura.",
)
replace_exact(document, "NPC com LLM generativo", "LLM gerando doutrina, profecia, revelação ou aconselhamento espiritual autoritativo")
replace_exact(
    document,
    "Davi é o primeiro vertical slice. A plataforma foi pensada para suportar novas jornadas, diferentes tipos de missão, AR quando fizer sentido, colecionáveis por história, evolução do Peregrino e uma mecânica de constância semanal.",
    "Davi é o primeiro vertical slice. A demonstração combina identidade visual forte, calibragem adaptativa, cinco minigames, orientação contextual segura, AR quando disponível, colecionáveis por história, evolução do Peregrino e constância semanal.",
)

# Atualiza os principais callouts e o backlog executivo.
replace_cell_exact(
    document,
    "OBJETIVO DO PILOTO\nProvar, em 36 horas, um loop completo e memorável: o jovem entra, cria seu Peregrino, escolhe uma jornada bíblica, aprende jogando, vive uma missão, consolida o aprendizado, recebe recompensas e percebe sua evolução.",
    "OBJETIVO DO PILOTO\nProvar, em menos de 24 horas, um loop completo e memorável: o jovem entra, cria seu Peregrino, responde a um diálogo de calibragem, recebe uma trilha adaptativa, aprende jogando, vive cinco etapas, enfrenta o vale, recebe recompensas e percebe sua evolução.",
)
replace_cell_exact(
    document,
    "LOOP DE SESSÃO\nEntrar → criar Peregrino → escolher Jornada → explorar a narrativa jogando → cumprir missão → consolidar aprendizado → receber recompensas → evoluir Peregrino → manter constância → voltar amanhã.",
    "LOOP DE SESSÃO\nEntrar → criar Peregrino → conversar e calibrar perfil → abrir mapa → explorar a narrativa jogando → consultar Parakletos quando necessário → cumprir missão → consolidar aprendizado → receber recompensas → evoluir → manter constância → voltar.",
)
replace_cell_exact(
    document,
    "DAVI E GOLIAS\nDescoberta → Prepare Davi → Missão AR: enfrente Golias → Quiz → Recompensas → Evolução do Peregrino.",
    "DAVI - VERTICAL SLICE\nAlém da Aparência → Guardião do Campo → A Missão → O que Você Conhece → Vale de Elá em AR ou simulação → Consolidação → Recompensas → Evolução do Peregrino.",
)

backlog = document.tables[9]
new_backlog_rows = [
    ("US18", "Calibragem imersiva do perfil", "P0", "Front + IA + PO"),
    ("US19", "Gerar trilha adaptativa com fallback", "P0", "Arquitetura/IA"),
    ("US20", "Consultar Parakletos com segurança", "P0", "Front + IA + PO"),
    ("US21", "Jogar cinco minigames com tempo", "P0", "Front + PO"),
    ("US22", "Recuperar vidas sem bloquear a jornada", "P0", "Front + Arquitetura"),
    ("US23", "Progredir na escada de materiais", "P0", "Front + Arquitetura"),
    ("US24", "Usar AR real ou fallback jogável", "P0 CRÍTICO", "VC/AR + Front"),
]
for values in new_backlog_rows:
    cells = backlog.add_row().cells
    for index, value in enumerate(values):
        cells[index].text = value

# Reescreve o Definition of Done existente sem aumentar o bloco visual.
dod_old = [
    "1.  Abrir",
    "2.  Iniciar Jornada",
    "3.  Criar Peregrino",
    "4.  Jornadas: Davi disponível; demais bloqueadas",
    "5.  Davi — Fé e Coragem",
    "6.  A Provação",
    "7.  Minigame: descobrir cenário",
    "8.  Minigame: preparar Davi",
    "9.  AR: Davi x Golias",
    "10.  Quiz de consolidação",
    "11.  Missão concluída",
    "12.  Receber colecionável",
    "13.  Atualizar constância",
    "14.  Evoluir equipamento do Peregrino",
    "15.  Ver Estante + Inventário",
]
dod = [
    "1. Abrir a experiência em tela cheia e iniciar sem cadastro",
    "2. Criar ou escolher um Peregrino com preview imediato",
    "3. Responder às cinco perguntas do diálogo de calibragem",
    "4. Receber um perfil e uma trilha adaptativa explicável",
    "5. Abrir o mapa da Jornada de Davi com cinco etapas",
    "6. Concluir Além da Aparência por investigação de pistas",
    "7. Concluir Guardião do Campo por ação e reflexo",
    "8. Concluir A Missão por memória e sequência",
    "9. Concluir O que Você Conhece escolhendo os itens coerentes",
    "10. Enfrentar o Vale de Elá por WebXR ou simulação de precisão",
    "11. Consultar Parakletos em qualquer etapa sem quebrar o fluxo",
    "12. Receber feedback, perder/recuperar vidas e continuar jogando",
    "13. Receber XP, material, colecionável e evolução de equipamento",
    "14. Atualizar constância de 3/7 para 4/7 e visualizar a meta",
    "15. Recarregar, abrir Estante/Inventário e reiniciar a demonstração",
]
for old, text in zip(dod_old, dod):
    replace_exact(document, old, text)

# Extensão consolidada: novas decisões e histórias, preservando o corpo original.
document.add_page_break()
document.add_heading("14. Decisões Consolidadas do Protótipo", level=1)

document.add_heading("14.1 Princípio de experiência", level=2)
document.add_paragraph(
    "O protótipo deve ser percebido como uma jornada jogável, não como uma sequência de formulários ou dashboards. A tela é ocupada pelo mundo, pelo personagem e pelo desafio; controles aparecem no contexto e usam áreas de toque grandes. A referência competitiva é a energia e o apelo de jogos juvenis contemporâneos, com identidade própria e sem copiar interfaces, marcas ou ativos de terceiros."
)
for text in [
    "Mobile-first e tela cheia, com HUD contínuo de vidas, tempo, XP e acesso ao guia.",
    "Menos botões de navegação; mais objetos, caminhos, mapas e elementos da própria cena como interação.",
    "Pixel art como linguagem visual unificadora: peregrinos, Parakletos, itens, materiais e evolução de equipamento.",
    "A interface deve continuar compreensível por uma pessoa que nunca recebeu explicação do time.",
]:
    add_bullet(document, text)

document.add_heading("14.2 Diálogo de calibragem e personalização", level=2)
document.add_paragraph(
    "A personalização começa com um diálogo imersivo conduzido como conversa narrativa, e não como formulário. O objetivo é ajustar linguagem, ritmo e mecânica - nunca medir a qualidade da fé do jogador."
)
for text in [
    "Faixa etária, sem solicitar data de nascimento.",
    "Contato prévio com histórias cristãs em casa, sempre com opção de não responder.",
    "Relação atual com a fé cristã, sempre opcional e sem classificação espiritual.",
    "Preferência entre ação/reflexo, puzzle/lógica, narrativa/leitura ou equilíbrio.",
    "Preferência narrativa por heróis improváveis, aventura/descoberta ou relações/escolhas.",
]:
    add_bullet(document, text)
document.add_paragraph(
    "A saída é um perfil de experiência - Guerreiro, Estrategista, Sábio ou Explorador - e um modo textual, gamificado ou equilibrado. A cronologia bíblica permanece preservada; a adaptação altera linguagem, ritmo, pistas, tempo e variação dos minigames. Dados sensíveis permanecem locais. Quando houver chamada de LLM, apenas o perfil derivado necessário é enviado."
)

document.add_heading("14.3 Uso de LLM e fallback", level=2)
for text in [
    "A LLM atua como diretora de experiência: escolhe variações disponíveis, ajusta dificuldade e redige orientação curta dentro de conteúdo aprovado.",
    "A resposta é estruturada e validada contra IDs conhecidos; a aplicação nunca executa nomes de cenas ou comandos arbitrários retornados pelo modelo.",
    "O endpoint é server-side, com chave fora do navegador, limites de tamanho e fallback determinístico local.",
    "Falha, latência ou ausência de credenciais não pode bloquear a jornada. O mesmo fluxo deve continuar com regras locais.",
    "A LLM não interpreta a vontade de Deus, não profetiza, não substitui aconselhamento pastoral e não avalia a condição espiritual do jogador.",
]:
    add_bullet(document, text)

document.add_heading("14.4 Parakletos - guia contextual", level=2)
document.add_paragraph(
    "Parakletos é uma pomba-personagem presente ao longo da jornada como guia educacional sempre disponível. Ele oferece contexto histórico, esclarece regras, aponta referências e ajuda o jogador a retomar o raciocínio."
)
for text in [
    "Deve ser apresentado explicitamente como personagem-guia e símbolo visual.",
    "Não é Deus, não é o Espírito Santo e não fala em nome de Deus.",
    "Não transmite profecias, revelações, absolvição, condenação ou aconselhamento espiritual autoritativo.",
    "Toda orientação bíblica vem acompanhada de referência e usa paráfrases próprias revisadas pelo PO.",
    "Pode oferecer pistas graduais, mas não entrega automaticamente a resposta do minigame.",
]:
    add_bullet(document, text)

document.add_heading("14.5 Cinco etapas jogáveis de Davi", level=2)
stages = [
    ("Além da Aparência", "Investigação de pistas", "1 Samuel 16:6-13", "Distinguir primeira impressão de caráter e perceber que Davi estava fora da seleção inicial."),
    ("Guardião do Campo", "Ação e reflexo", "1 Samuel 17:34-37", "Proteger o rebanho e conectar responsabilidade cotidiana com preparo."),
    ("A Missão", "Memória e sequência", "1 Samuel 17:17-22", "Reconstruir a entrega de mantimentos e perceber que a virada começa em uma tarefa comum."),
    ("O que Você Conhece", "Seleção de equipamento", "1 Samuel 17:38-40", "Escolher cajado, funda e cinco pedras; armadura, espada e escudo funcionam como distratores com feedback."),
    ("Vale de Elá", "Precisão e AR", "1 Samuel 17:41-51", "Posicionar a cena e lançar a pedra em WebXR ou concluir o fallback de ritmo e precisão."),
]
for title, mechanic, reference, purpose in stages:
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{title} - {mechanic} | {reference}\n").bold = True
    paragraph.add_run(purpose)
    paragraph.paragraph_format.space_after = Pt(6)

document.add_heading("14.6 Tempo, vidas e continuidade", level=2)
for text in [
    "Cada etapa possui tempo ajustado ao modo textual, gamificado ou equilibrado.",
    "O jogador começa com três vidas; erro e tempo esgotado geram feedback específico e nova tentativa.",
    "Zerar vidas não encerra nem monetiza a experiência. Abre uma tarefa de recuperação ligada à passagem bíblica em que ocorreu a dificuldade.",
    "Na versão de produção, a recuperação pode oferecer leitura e lembrete em uma hora. No hackathon, a retomada é imediata para preservar a demonstração.",
    "A dificuldade deve valorizar atenção, compreensão e habilidade, sem transformar conhecimento bíblico prévio em vantagem injusta.",
]:
    add_bullet(document, text)

document.add_heading("14.7 XP, materiais e recompensas", level=2)
document.add_paragraph(
    "Tempo restante, precisão, tentativas e conclusão alimentam uma escada de materiais visual. Ela comunica desempenho sem substituir os colecionáveis narrativos ou o progresso de constância."
)
for text in [
    "5 XP - Palha",
    "15 XP - Capim",
    "25 XP - Madeira",
    "50 XP - Prata",
    "80 XP - Ouro",
    "110 XP - Pedras preciosas",
]:
    add_bullet(document, text)
document.add_paragraph(
    "A Pedra de Davi é concedida pela conclusão. O Escudo da Fé evolui conforme score: comum até 59, raro de 60 a 84 e épico a partir de 85. A Estante guarda memórias das histórias; o Inventário guarda materiais e equipamentos do Peregrino."
)

document.add_heading("14.8 Stack e contratos do protótipo", level=2)
for text in [
    "React + TypeScript + Vinext/Vite; CSS Modules e variáveis CSS; experiência sem biblioteca visual pesada.",
    "Zustand com persistência local, schema versionado e botão discreto de reinício.",
    "Conteúdo bíblico, perfis e definições de minigames mantidos em estruturas locais revisáveis.",
    "Endpoint server-side compatível com OpenAI Responses API e fallback curado.",
    "Three.js carregado apenas quando necessário e adaptador isolado para WebXR immersive-ar, hit-test e referência local.",
    "Localhost como ambiente de desenvolvimento; Android Chrome + ARCore e contexto seguro para a demonstração real de AR.",
]:
    add_bullet(document, text)

document.add_heading("14.9 Definition of Done atualizado", level=2)
for text in [
    "Três demonstrações consecutivas após reset, sem intervenção técnica.",
    "Uma pessoa externa conclui o fluxo em até cinco minutos.",
    "Progresso, recompensas, perfil e constância sobrevivem ao recarregamento.",
    "O confronto final é jogável no desktop e oferece WebXR real em Android compatível.",
    "Parakletos permanece acessível e exibe seu posicionamento de personagem-guia.",
    "Typecheck, testes, build de produção e abertura local são executados antes do congelamento da demo.",
]:
    add_bullet(document, text)

document.add_page_break()
document.add_heading("15. Novas User Stories Priorizadas", level=1)
add_story(
    document,
    "US18",
    "Concluir uma calibragem imersiva",
    "Como novo Peregrino, quero responder a um diálogo curto e narrativo, para receber uma experiência adequada ao meu ritmo sem preencher um formulário convencional.",
    [
        "Apresentar cinco perguntas em linguagem de jornada, uma por vez.",
        "Permitir não responder perguntas relacionadas a fé ou ambiente familiar.",
        "Explicar que as respostas ajustam experiência e não medem espiritualidade.",
        "Concluir a calibragem em aproximadamente um minuto.",
    ],
    "Front + IA + PO",
)
add_story(
    document,
    "US19",
    "Receber uma trilha adaptativa",
    "Como Peregrino, quero que a jornada ajuste ritmo, pistas e linguagem às minhas escolhas, para sentir que a experiência foi preparada para mim.",
    [
        "Gerar um dos quatro perfis explicáveis: Guerreiro, Estrategista, Sábio ou Explorador.",
        "Selecionar modo textual, gamificado ou equilibrado sem alterar fatos ou cronologia bíblica.",
        "Usar somente IDs de minigames previamente aprovados.",
        "Usar fallback local determinístico se a LLM falhar ou estiver indisponível.",
    ],
    "Arquitetura/IA + Front",
)
add_story(
    document,
    "US20",
    "Consultar Parakletos durante a jornada",
    "Como Peregrino, quero acessar um guia contextual sem sair da etapa atual, para compreender o contexto e receber pistas quando necessário.",
    [
        "Manter acesso ao guia no HUD e nos momentos relevantes.",
        "Retornar contexto curto, pista gradual e referência bíblica.",
        "Exibir o posicionamento de personagem-guia e os limites teológicos.",
        "Nunca usar linguagem de revelação, profecia ou autoridade divina.",
    ],
    "Front + IA + PO",
)
add_story(
    document,
    "US21",
    "Jogar cinco minigames contextualizados",
    "Como Peregrino, quero enfrentar desafios variados e dinâmicos, para aprender detalhes reais da história sem sentir que estou respondendo a uma prova escolar.",
    [
        "Oferecer investigação, reflexo, memória, seleção e precisão/AR.",
        "Associar cada etapa a contexto histórico, pergunta provocativa e referência.",
        "Exibir tempo e feedback imediato sem bloquear a conclusão por erro.",
        "Adaptar ritmo e quantidade de texto ao perfil de experiência.",
    ],
    "Front + PO",
)
add_story(
    document,
    "US22",
    "Recuperar vidas sem abandonar a experiência",
    "Como Peregrino, quero ter uma forma significativa de continuar após perder minhas vidas, para que a dificuldade gere aprendizagem em vez de abandono.",
    [
        "Iniciar a jornada com três vidas.",
        "Ao zerar vidas, indicar uma leitura curta relacionada à etapa.",
        "Restaurar uma vida após a tarefa de recuperação.",
        "Na demo, permitir retomada imediata; em produção, suportar lembrete opcional em uma hora.",
    ],
    "Front + Arquitetura",
)
add_story(
    document,
    "US23",
    "Receber materiais conforme desempenho",
    "Como Peregrino, quero visualizar a qualidade da minha execução por materiais progressivos, para ter um motivo claro para melhorar sem perder a recompensa narrativa.",
    [
        "Calcular XP a partir de conclusão, tempo e precisão.",
        "Mapear XP para palha, capim, madeira, prata, ouro e pedras preciosas.",
        "Adicionar materiais conquistados ao Inventário.",
        "Conceder o colecionável da história independentemente do material alcançado.",
    ],
    "Front + Arquitetura",
)
add_story(
    document,
    "US24",
    "Concluir o confronto em AR ou fallback",
    "Como Peregrino, quero enfrentar Golias no meu ambiente quando meu aparelho suportar AR, sem perder o capítulo quando estiver em desktop ou aparelho incompatível.",
    [
        "Detectar suporte a immersive-ar antes de iniciar a sessão.",
        "No Android compatível, detectar superfície, posicionar a cena, lançar a pedra e reagir ao acerto.",
        "Em ambiente incompatível, oferecer minigame de ritmo e precisão com a mesma conclusão.",
        "Tratar permissão negada e saída de sessão sem perder progresso.",
    ],
    "VC/AR + Front",
)

# Ajustes finais de legibilidade nos novos conteúdos.
for paragraph in document.paragraphs:
    if paragraph.style.name.startswith("Heading"):
        keep_with_next(paragraph)
    if paragraph.style.name == "Normal":
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
document.save(OUTPUT)
print(OUTPUT)
